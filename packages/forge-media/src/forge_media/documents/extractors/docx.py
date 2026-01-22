"""DOCX document extraction using python-docx."""

from __future__ import annotations

from io import BytesIO

from forge_media.documents.extractors.base import (
    DocumentExtractor,
    DocumentPage,
    ExtractedDocument,
    ExtractedTable,
    ExtractionError,
    ExtractionOptions,
)


class DOCXExtractor(DocumentExtractor):
    """
    Extract content from Microsoft Word (.docx) documents.

    Uses python-docx for extraction. Note that DOCX files don't have
    explicit pages, so all content is returned as a single page.

    Example:
        ```python
        extractor = DOCXExtractor()
        with open("document.docx", "rb") as f:
            data = f.read()

        result = await extractor.extract(data)
        print(result.full_text)
        for table in result.tables:
            print(f"Table with {table.row_count} rows")
        ```
    """

    @property
    def supported_mime_types(self) -> list[str]:
        return [
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/msword",  # Legacy .doc format (limited support)
        ]

    async def extract(
        self,
        data: bytes,
        options: ExtractionOptions | None = None,
    ) -> ExtractedDocument:
        """
        Extract content from a DOCX document.

        Args:
            data: Raw DOCX bytes
            options: Extraction options

        Returns:
            ExtractedDocument with content (single page for DOCX)
        """
        options = options or ExtractionOptions()

        try:
            from docx import Document
        except ImportError as e:
            raise ExtractionError(
                "python-docx is required for DOCX extraction. "
                "Install with: pip install python-docx",
                cause=e,
            )

        try:
            doc = Document(BytesIO(data))
        except Exception as e:
            raise ExtractionError(f"Failed to read DOCX: {e}", cause=e)

        # Extract paragraphs
        text_parts: list[str] = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        # Extract tables
        tables: list[ExtractedTable] = []
        if options.extract_tables:
            for table_idx, table in enumerate(doc.tables):
                extracted = self._extract_table(table, table_idx)
                if extracted.rows:  # Only add non-empty tables
                    tables.append(extracted)

        full_text = "\n".join(text_parts)

        # DOCX doesn't have pages, create single page with all content
        page = DocumentPage(
            page_number=1,
            text=full_text,
            tables=tables,
            images=[],
        )

        # Extract metadata
        metadata = {}
        if options.extract_metadata:
            props = doc.core_properties
            metadata = {
                "title": props.title,
                "author": props.author,
                "subject": props.subject,
                "keywords": props.keywords,
                "created": props.created.isoformat() if props.created else None,
                "modified": props.modified.isoformat() if props.modified else None,
                "last_modified_by": props.last_modified_by,
            }
            # Remove None values
            metadata = {k: v for k, v in metadata.items() if v is not None}

        return ExtractedDocument(
            pages=[page],
            full_text=full_text,
            tables=tables,
            metadata=metadata,
            page_count=1,
            word_count=len(full_text.split()),
        )

    def _extract_table(self, table, index: int) -> ExtractedTable:
        """Extract data from a DOCX table."""
        rows: list[list[str]] = []

        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            rows.append(row_data)

        # Use first row as headers if table has data
        headers = rows[0] if rows else None
        data_rows = rows[1:] if len(rows) > 1 else []

        return ExtractedTable(
            page_number=1,  # DOCX doesn't have pages
            table_index=index,
            headers=headers,
            rows=data_rows,
        )
